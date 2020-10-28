from __future__ import print_function
from docx import Document
from docx.shared import Pt,Cm,Inches
from mailmerge import MailMerge
import datetime
from docx2pdf import convert
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
from pdfCvt import convert_to
import inflect
import sys
import json

def cellVerticalAlignment(table):
    for row in table.rows:
        for cell in row.cells:
            pass
            #cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_col_widths(table):
    widths = (Inches(1),Inches(0.551181),Inches(1.5),Inches(0.6),Inches(0.9),Inches(1.5))
    table.columns[0].width = Inches(1)
    table.columns[1].width = Inches(0.551181)
    table.columns[2].width = Inches(1.5)
    table.columns[3].width = Inches(0.6)
    table.columns[4].width = Inches(0.9)
    table.columns[5].width = Inches(1.5)

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def modifyBorder(table):
    tbl = table._tbl # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'nil')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'nil')

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

def preventDocumentBreak(document):
  tags = document.element.xpath('//w:tr')
  rows = len(tags)
  for row in range(0,rows):
    tag = tags[row]
    child = OxmlElement('w:cantSplit')
    tag.append(child)

def changeSize(cell,size):
    paragraph =cell.paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(size)



def createPdfContractor(path):
    path = str(path).split("/")
    folder = path[0]
    filename = path[1]
    items = {}
    path = "{}\\{}.csai".format(str(folder),str(filename))


    if os.path.isfile(path):
        with open(path, 'r') as fp:
            items=json.load(fp)
    contractor_name = items['contractor_name']
    phone_no = items['phone_no']
    team_of = items['team_of']
    rate = items['rate']
    month = items['month']
    year = items['year']
    template = "templates\\contractor_data.docx"
    document = MailMerge(template)

    document.merge(contractor_name=contractor_name,rate=rate, phone_no=phone_no, team_of=team_of, month=month, year=year)

    document.write("temp.docx")

    doc = Document(docx='temp.docx')
    table = doc.tables[1]
    cell = table.rows[0].cells[0]
    #cell.text = item['item_name']
    inner_table = cell.add_table(rows=1,cols=6)
    inner_hdr_cell = inner_table.rows[0].cells
    inner_hdr_cell[0].text = 'Date'
    inner_hdr_cell[1].text = 'Weight'
    inner_hdr_cell[2].text = 'Customer Name'
    inner_hdr_cell[3].text = 'Advance'
    inner_hdr_cell[4].text = 'By Hand'
    inner_hdr_cell[5].text = 'Description'

    for i in range(0,6):
        inner_hdr_cell[i].paragraphs[0].runs[0].font.bold = True

    for j in range(int(items['total_rows'])):
        inner_celss = inner_table.add_row().cells
        if j in items.keys():
            bundle = items[j]
        else:
            bundle = items[str(j)]
        inner_celss[0].text = str(bundle['date']).strip()
        if str(bundle['weight'])=='' or str(bundle['weight'])=='-':
            inner_celss[1].text = '-'
        else:
            inner_celss[1].text = str(bundle['weight'])
        if str(bundle['cust_name'])=='' or str(bundle['cust_name'])=='-':
            inner_celss[2].text = '-'
        else:
            inner_celss[2].text = str(bundle['cust_name']).strip()

        if str(bundle['advance'])=='' or str(bundle['advance'])=='-':
            inner_celss[3].text = '-'
        else:
            inner_celss[3].text = str(bundle['advance']).strip()
        if str(bundle['by_hand'])=='' or str(bundle['by_hand'])=='-':
            inner_celss[4].text = '-'
        else:
            inner_celss[4].text = str(bundle['by_hand']).strip()
        if str(bundle['description'])=='' or str(bundle['description'])=='-':
            inner_celss[5].text = '-'
        else:
            inner_celss[5].text = str(bundle['description']).strip()

    inner_table.style = "Table Grid"
    set_col_widths(inner_table)
    table = doc.tables[2]
    cell = table.rows[0].cells
    tot_bill = 0.0
    try:
        tot_bill = float(items['total_weight'])*float(rate)
    except:
        tot_bill = 0.0
    cell[1].text = " {} KG".format(items['total_weight'])
    cell[3].text = " {} ".format(tot_bill)
    paragraph =cell[1].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    paragraph =cell[3].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    cell = table.rows[1].cells
    cell[1].text = " {} ".format(items['total_advance'])
    if tot_bill>0:
        try:
            tot_bill = tot_bill-float(items['total_advance'])
        except:
            pass
    cell[3].text = " {} ".format(tot_bill)

    paragraph =cell[1].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    paragraph =cell[3].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True


    doc.save('temp.docx')
    convert_to("temp.docx")
    os.remove("temp.docx")
    #os.system(f"pdf_contractor\\{filename}.pdf")

if __name__ == "__main__":
    createPdfContractor(*sys.argv[1:])