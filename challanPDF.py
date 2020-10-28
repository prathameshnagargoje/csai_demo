from __future__ import print_function
from docx import Document
from docx.shared import Pt,Cm,Inches
from mailmerge import MailMerge
import datetime
from docx2pdf import convert
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
from pdfCvt import convert_to
import inflect
import sys
import json

def cellVerticalAlignment(table):
    for row in table.rows:
        for cell in row.cells:
            pass
            #cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_col_widths(table):
    widths = (Inches(0.459318),Inches(0.498688),Inches(0.734908))
    table.columns[0].width = Inches(0.459318)
    table.columns[1].width = Inches(0.498688)
    table.columns[2].width = Inches(0.734908)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def modifyBorder(table):
    tbl = table._tbl # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'nil')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'nil')

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

def preventDocumentBreak(document):
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0,rows):
        tag = tags[row]
        child = OxmlElement('w:cantSplit')
        tag.append(child)

def changeSize(cell,size,b='n'):
    paragraph =cell.paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(size)
    if b=='y':
        font.bold = True

def createPdfChallan(path):
    path = str(path).split("/")
    folder = path[0]
    filename = path[1]
    items = {}
    path = "{}\\{}.csai".format(str(folder),str(filename))

    if os.path.isfile(path):
        with open(path, 'r') as fp:
            items=json.load(fp)
    
    cust_name = items['cust_name'].strip()
    dday = datetime.datetime.now()
    date_day = dday.strftime("%d-%b-%y, %A")
    challan_no = str(items['challan_no']).strip()
    vehicle_no = items['vehicle_no'].strip()
    total_weight = str(items['total_weight']).strip()
    template = "templates\\challan.docx"
    document = MailMerge(template)

    document.merge(cust_name=cust_name, date_day=date_day, challan_no=challan_no, vehicle_no=vehicle_no,
     total_weight=total_weight)

    document.write("temp.docx")

    doc = Document(docx='temp.docx')
    table = doc.tables[0]
    row = table.add_row()
    cell = row.cells[0]
    #cell = table.rows[0].cells[0]

    for i in range(0,items['tot_items']):
        if i in items.keys():
            item = items[i]
        elif str(i) in items.keys():
            item = items[str(i)]
        else:
            continue
        inner_table = cell.add_table(rows=1,cols=3)
        a,b,c = inner_table.rows[0].cells
        a.text = item['item_name'].strip()
        a.merge(b).merge(c)
        changeSize(a,10,'y')
        row = inner_table.add_row()
        
        inner_hdr_cell = row.cells
        inner_hdr_cell[0].text = 'No.'
        paragraph =inner_hdr_cell[0].paragraphs[0]
        paragraph.alignment = 1
    
        inner_hdr_cell[1].text = 'Qty'
        paragraph =inner_hdr_cell[1].paragraphs[0]
        paragraph.alignment = 1
    
        inner_hdr_cell[2].text = 'Weight'
        paragraph =inner_hdr_cell[2].paragraphs[0]
        paragraph.alignment = 1

        changeSize(inner_hdr_cell[0],8)
        changeSize(inner_hdr_cell[1],8)
        changeSize(inner_hdr_cell[2],8)

        for j in range(0,item['tot_bundles']):
            if j in item.keys():
                bundle = item[j]
            elif str(j) in item.keys():
                bundle = item[str(j)]
            else:
                continue
            row = inner_table.add_row()
            inner_celss = row.cells
            
            inner_celss[0].text = str(j+1).strip()
        
            try:
                inner_celss[1].text = "{:.0f}".format(float(str(bundle['qty']).strip()))
            except:
                inner_celss[1].text = str(bundle['qty']).strip().split('.')[0]
        
            inner_celss[2].text = str(bundle['gross_wt']).strip()
            
            paragraph =inner_celss[0].paragraphs[0]
            paragraph.alignment = 1
            paragraph =inner_celss[1].paragraphs[0]
            paragraph.alignment = 1
            paragraph =inner_celss[2].paragraphs[0]
            paragraph.alignment = 1
            changeSize(inner_celss[0],8)
            changeSize(inner_celss[1],8)
            changeSize(inner_celss[2],8)
        row = inner_table.add_row()
        total_cells = row.cells
        total_cells[0].text = " Total "
        changeSize(total_cells[0],8,'y')
        row = inner_table.add_row()

        total_cells = row.cells
    
        total_cells[0].text = f"{str(item['tot_bundles']).strip()}"
        paragraph =total_cells[0].paragraphs[0]
        paragraph.alignment = 1
        total_cells[1].text = str(round(float(str(item['tot_qty']).strip())))
        paragraph =total_cells[1].paragraphs[0]
        paragraph.alignment = 1
        total_cells[2].text = str(item['tot_gross_wt']).strip()
        paragraph =total_cells[2].paragraphs[0]
        paragraph.alignment = 1
        changeSize(total_cells[0],8,'y')
        changeSize(total_cells[1],8,'y')
        changeSize(total_cells[2],8,'y')
        inner_table.style = "Table Grid"
        set_col_widths(inner_table)
        row = table.add_row()
        cell = row.cells[0]

    preventDocumentBreak(doc)
    doc.save('temp.docx')
    #if not os.path.isdir("pdf_challan"):
    #os.makedirs("pdf_challan")
    convert_to("temp.docx")
    os.remove("temp.docx")
    #os.system("pdf_challan\\{filename}.pdf")

if __name__ == "__main__":
    createPdfChallan(*sys.argv[1:])