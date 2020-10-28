from __future__ import print_function
from docx import Document
from docx.shared import Pt,Cm,Inches
from mailmerge import MailMerge
import datetime
from docx2pdf import convert
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
from pdfCvt import convert_to
import inflect
import sys
import json

def cellVerticalAlignment(table):
    for row in table.rows:
        for cell in row.cells:
            pass
            #cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_col_widths(table):
    widths = (Inches(0.275591),Inches(0.551181),Inches(0.314961),Inches(0.551181))
    table.columns[0].width = Inches(0.275591)
    table.columns[1].width = Inches(0.551181)
    table.columns[2].width = Inches(0.314961)
    table.columns[3].width = Inches(0.551181)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def modifyBorder(table):
    tbl = table._tbl # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'nil')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'nil')

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

def preventDocumentBreak(document):
  tags = document.element.xpath('//w:tr')
  rows = len(tags)
  for row in range(0,rows):
    tag = tags[row]
    child = OxmlElement('w:cantSplit')
    tag.append(child)

def changeSize(cell,size):
    paragraph =cell.paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(size)



def createPdfBill(path):

    path = str(path).split("/")
    folder = path[0]
    filename = path[1]
    items = {}
    path = "{}\\{}.csai".format(str(folder),str(filename))
    if os.path.isfile(path):
        with open(path, 'r') as fp:
            items=json.load(fp)
    
    invoice_no = str(items['invoice_no'])
    invoice_date = str(items['invoice_date'])
    supply_date = str(items['supply_date'])
    transport = str(items['transport'])
    vehicle_no = str(items['vehicle_no'])
    cust_name = str(items['cust_name'])
    cust_address = str(items['cust_address'])
    cust_state = str(items['cust_state'])
    cust_gstin = str(items['cust_gstin'])
    cgst_per = str(items['cgst_per'])
    sgst_per = str(items['sgst_per'])
    igst_per = str(items['igst_per'])
    try:
        cgst_total = "{:.3f}".format(float(items['cgst_total']))
        sgst_total = "{:.3f}".format(float(items['sgst_total']))
        igst_total = "{:.3f}".format(float(items['igst_total']))
        grand_total = "{:.3f}".format(float(items['grand_total']))
    except:
        pass
    destination = str(items['destination'])
    state_code = str(items['state_code'])
    if str(items['if_copy']).strip() == 'F':
        if_copy = "(Original For Recipient)"
    else:
        if_copy = "(Duplicate Copy)"

    amount = str(inflect.engine().number_to_words(float(grand_total))).split("point")
    amount_word = str(amount[0]).strip().capitalize()+" rupees and "+amount[1]+" paisa only"
    counter = int(items['counter'])

    template = "templates\\bill.docx"
    document = MailMerge(template)

    document.merge(invoice_no=invoice_no,date_invoice=invoice_date,date_supply=supply_date, transport=transport,
                        vehicle_no=vehicle_no, cust_name=cust_name, cust_address=cust_address, cust_state=cust_state,
                        cust_gstin=cust_gstin, cgst_per=cgst_per, sgst_per=sgst_per,igst_per=igst_per,
                         cgst_total=cgst_total, sgst_total=sgst_total,igst_total=igst_total,
                        grand_total=grand_total, amount_word=amount_word, if_copy=if_copy,
                         destination=destination, state_code=state_code)

    document.write("temp.docx")

    doc = Document(docx='temp.docx')

    table = doc.tables[2]
    data_cells = table.rows[1].cells

    tot_qty = 0
    tot_weight = 0.0
    tot_amount = 0.0

    for i in range(0,counter):
        if i in items.keys():
            row_data = items[i]
        else:
            row_data = items[str(i)]
        data_cells[0].text = str(i+1).strip()
        data_cells[1].text = str(row_data['description']).strip()
        data_cells[2].text = str(row_data['hsncode']).strip()
        data_cells[3].text = str(row_data['qty']).strip()
        data_cells[4].text = str(row_data['weight']).strip()
        data_cells[5].text = str(row_data['price']).strip()
        data_cells[6].text = str(row_data['amount']).strip()
        try:
            tot_qty+=int(str(row_data['qty']).strip())
        except:
            pass
        try:
            tot_weight+=float(str(row_data['weight']).strip())
        except:
            pass
        try:
            tot_amount+=float(str(row_data['amount']).strip())
        except:
            pass
        
        if (i+1)<counter:
            data_cells = table.add_row().cells

    data_cells = table.add_row().cells
    data_cells[0].text = "Sub Total: "
    data_cells[0].merge(data_cells[1]).merge(data_cells[2])
    paragraph =data_cells[0].paragraphs[0]
    paragraph.alignment = 2
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    data_cells[3].text = str(tot_qty)
    paragraph =data_cells[3].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    data_cells[4].text = str(tot_weight)
    paragraph =data_cells[4].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True
    data_cells[6].text = str(tot_amount)
    paragraph =data_cells[6].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.bold = True


    doc.save('temp.docx')
    convert_to("temp.docx")
    os.remove("temp.docx")
    #os.system(f"pdf_bill\\{filename}.pdf")

if __name__ == "__main__":
    createPdfBill(*sys.argv[1:])